# =============================================
# procesar_anexo5.py
# =============================================
import os
import re
import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
import yaml

# ============================================================
# ‚öôÔ∏è CONFIGURACI√ìN DE RUTAS
# ============================================================
BASE_DIR = Path(__file__).resolve().parents[1]
CONFIG_PATH = BASE_DIR / "config/settings_anexo5.yaml"
CARPETA_RAW = BASE_DIR / "data/raw"

# ============================================================
# üìÑ CARGAR CONFIGURACI√ìN YAML
# ============================================================
if CONFIG_PATH.exists():
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        CONFIG = yaml.safe_load(f)
else:
    CONFIG = {
        "salida": {
            "carpeta": "data/processed/anexo5",
            "archivo_excel": "anexo5_consolidado.xlsx"
        }
    }

# ============================================================
# üß© FUNCIONES AUXILIARES Y LOG
# ============================================================
def log(mensaje, tipo="INFO", anexo="ANEXO_5"):
    hora = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    simbolo = {"INFO": "‚ÑπÔ∏è", "OK": "‚úÖ", "WARN": "‚ö†Ô∏è", "ERROR": "‚ùå"}.get(tipo, "")
    print(f"{simbolo} [{anexo}] [{hora}] {mensaje}")


def leer_docx(path):
    """Lee el texto completo y las tablas del documento Word."""
    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    return full_text, doc.tables


def extraer_metadatos(text, doc_path):
    """Extrae Unidad Territorial y Fecha del documento o carpeta."""
    clean_text = re.sub(r"[\r\t\xa0]+", " ", text)
    clean_text = re.sub(r"\s+", " ", clean_text)

    unidad_match = re.search(r"UNIDAD\s*TERRITORIAL\s*:\s*([A-Z√Å√â√ç√ì√ö√ë ]+)", clean_text, re.IGNORECASE)
    fecha_match = re.search(r"FECHA\s*:\s*([\d/]+)", clean_text, re.IGNORECASE)

    if unidad_match:
        unidad = unidad_match.group(1).strip().upper()
    else:
        unidad = os.path.basename(os.path.dirname(doc_path)).replace("_", " ").upper()

    fecha = fecha_match.group(1).strip() if fecha_match else None

    return {"unidad_territorial": unidad, "fecha": fecha}


def leer_tabla_docx(tables):
    """Extrae todas las filas de las tablas del documento."""
    data = []
    for table in tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                data.append(cells)
    return data


def procesar_tabla(raw_data, meta, region, mes, anio, archivo):
    """Convierte la tabla de Word en un DataFrame."""
    filas = [f for f in raw_data if f and f[0].strip().isdigit()]

    registros = []
    for fila in filas:
        fila += [""] * (5 - len(fila))
        registros.append({
            "A√±o": anio,
            "Mes": mes,
            "Regi√≥n": region,
            "Archivo": archivo,
            "N¬∞": fila[0],
            "PUNTOS_CRITICOS": fila[1],
            "ACUERDOS_MEJORA": fila[2],
            "RESPONSABLE": fila[3],
            "PLAZO": fila[4],
            "UNIDAD_TERRITORIAL": meta["unidad_territorial"],
            "FECHA": meta["fecha"]
        })

    return pd.DataFrame(registros)


# ============================================================
# üöÄ PROCESAR TODOS LOS DOCUMENTOS
# ============================================================
if not CARPETA_RAW.exists():
    raise FileNotFoundError(f"‚ùå No existe la carpeta: {CARPETA_RAW}")

registros = []
for carpeta_anio in sorted(CARPETA_RAW.iterdir()):
    if not carpeta_anio.is_dir():
        continue
    anio = carpeta_anio.name
    for carpeta_mes in sorted(carpeta_anio.iterdir()):
        if not carpeta_mes.is_dir():
            continue
        mes = carpeta_mes.name
        log(f"Procesando: {anio}/{mes}", "INFO")
        for carpeta_region in sorted(carpeta_mes.iterdir()):
            if not carpeta_region.is_dir():
                continue
            region = carpeta_region.name.upper()
            archivos = list(carpeta_region.glob("*ANEXO_5*.docx"))
            if not archivos:
                log(f"{anio}/{mes}/{region}: sin archivos Word", "WARN")
                continue

            for archivo in archivos:
                try:
                    text, tables = leer_docx(archivo)
                    meta = extraer_metadatos(text, archivo)
                    raw_data = leer_tabla_docx(tables)
                    df = procesar_tabla(raw_data, meta, region, mes, anio, archivo.name)
                    registros.append(df)
                    log(f"{anio}/{mes}/{region}/{archivo.name} procesado", "OK")
                except Exception as e:
                    log(f"Error en {archivo}: {e}", "ERROR")

# ============================================================
# üíæ EXPORTAR CONSOLIDADO
# ============================================================
if registros:
    df_total = pd.concat(registros, ignore_index=True)

    salida_dir = BASE_DIR / CONFIG["salida"]["carpeta"]
    salida_dir.mkdir(parents=True, exist_ok=True)

    salida_excel = salida_dir / CONFIG["salida"]["archivo_excel"]

    if salida_excel.exists():
        df_prev = pd.read_excel(salida_excel)
        df_total = pd.concat([df_prev, df_total], ignore_index=True)
        df_total.drop_duplicates(subset=["Archivo", "Regi√≥n", "Mes", "A√±o"], inplace=True)

    df_total.to_excel(salida_excel, index=False)
    log(f"Consolidado actualizado: {salida_excel}", "OK")
    log(f"Total de registros acumulados: {len(df_total)}", "INFO")

else:
    log("No se proces√≥ ning√∫n documento.", "WARN")

log("Procesamiento de ANEXO 5 finalizado.", "OK")